from fastapi import FastAPI, UploadFile, File
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from transformers import AutoTokenizer, AutoModelForCausalLM

import os, json, time, pickle, shutil, re, io, contextlib, threading
from datetime import datetime

import torch
import numpy as np

try:
    import pandas as pd
except Exception:
    pd = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    import pypdfium2 as pdfium
except Exception:
    pdfium = None

try:
    import pytesseract
    from PIL import Image 
except Exception:
    pytesseract = None
    Image = None


try:
    import z3  
except Exception:
    z3 = None


if pytesseract is not None:
    _TES = os.getenv("TESSERACT_CMD")
    if _TES:
        pytesseract.pytesseract.tesseract_cmd = _TES

MODEL_ID = os.getenv("MODEL_ID", "meta-llama/Meta-Llama-3-8B-Instruct")
HF_TOKEN = os.getenv("HUGGINGFACE_HUB_TOKEN") or os.getenv("HF_TOKEN")
DEVICE = ("cuda" if torch.cuda.is_available()
          else ("mps" if torch.backends.mps.is_available() else "cpu"))

print("Loading model...")
tokenizer = AutoTokenizer.from_pretrained(MODEL_ID, token=HF_TOKEN, use_fast=True)
model = AutoModelForCausalLM.from_pretrained(
    MODEL_ID,
    torch_dtype=(torch.float16 if DEVICE == "cuda" else torch.float32),
    device_map="auto",
    token=HF_TOKEN,
)
device = next(model.parameters()).device
print(f"Model ready on device: {device}")


import faiss
from sentence_transformers import SentenceTransformer
from rank_bm25 import BM25Okapi

BASE_DIR   = os.path.dirname(__file__)
KNOW_DIR   = os.path.join(BASE_DIR, "knowledge")
os.makedirs(KNOW_DIR, exist_ok=True)

INDEX_FILE = os.path.join(BASE_DIR, "faiss_index.bin")
META_FILE  = os.path.join(BASE_DIR, "faiss_meta.pkl")

EMBED_MODEL       = "all-MiniLM-L6-v2"
TOP_K             = 8           
MAX_CONTEXT_CHARS = 2500


embedder: SentenceTransformer | None = None
faiss_index: faiss.Index | None = None
meta: list[dict] = []

bm25_corpus: list[list[str]] = []
bm25: BM25Okapi | None = None


SUPPORTED_EXTS = {".txt", ".md", ".docx", ".csv", ".pdf"}



def _tok(text: str) -> list[str]:
    return re.findall(r"[a-z0-9]+", (text or "").lower())


def _load_rag():
    global embedder, faiss_index, meta
    try:
        embedder = SentenceTransformer(EMBED_MODEL)
    except Exception as e:
        print(" Could not load embedder:", e)
        embedder = None

    try:
        faiss_index = faiss.read_index(INDEX_FILE) if os.path.exists(INDEX_FILE) else None
        if os.path.exists(META_FILE):
            with open(META_FILE, "rb") as f:
                meta = pickle.load(f)
        else:
            meta = []
        print(f"RAG ready: FAISS={'yes' if faiss_index is not None else 'no'}, docs={len(meta)}")
    except Exception as e:
        print(" Failed to load FAISS index/meta:", e)
        faiss_index = None
        meta = []


def _embed(texts):
    if embedder is None:
        return None
    return embedder.encode(texts, convert_to_numpy=True, show_progress_bar=False)


def _pdf_extract_text_pypdf(path: str, max_pages: int) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(path)
        return "\n".join((page.extract_text() or "") for page in reader.pages[:max_pages]).strip()
    except Exception:
        return ""


def _pdf_extract_text_pdfium(path: str, max_pages: int) -> str:
    if pdfium is None:
        return ""
    try:
        pdf = pdfium.PdfDocument(path)
        return "\n".join(pdf[i].get_textpage().get_text_range() or "" for i in range(min(max_pages, len(pdf)))).strip()
    except Exception:
        return ""


def _pdf_ocr_first_pages(path: str, pages_to_ocr: int = 2, scale: float = 2.0) -> str:
    if pdfium is None or pytesseract is None or Image is None:
        return ""
    try:
        pdf = pdfium.PdfDocument(path)
        texts = []
        for i in range(min(pages_to_ocr, len(pdf))):
            page = pdf[i]
            bitmap = page.render(scale=scale).to_pil()
            txt = pytesseract.image_to_string(bitmap)
            if txt:
                texts.append(txt)
        return "\n".join(texts).strip()
    except Exception:
        return ""


def _extract_text(path: str, for_snippet: bool = False) -> str:
    ext = os.path.splitext(path)[1].lower()
    txt = ""

    try:
        if ext in (".txt", ".md"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                txt = f.read()

        elif ext == ".docx":
            
            if Document is None:
                return f"[DOCX: {os.path.basename(path)}]"
            try:
                d = Document(path)
                parts = []

                for p in d.paragraphs:
                    if p.text.strip():
                        parts.append(p.text.strip())

                for tbl in d.tables:
                    for row in tbl.rows:
                        row_txt = " | ".join((c.text or "").strip() for c in row.cells)
                        if row_txt.strip():
                            parts.append(row_txt)

                for sec in d.sections:
                    for p in getattr(sec.header, "paragraphs", []):
                        if p.text.strip():
                            parts.append(p.text.strip())
                    for p in getattr(sec.footer, "paragraphs", []):
                        if p.text.strip():
                            parts.append(p.text.strip())

                txt = "\n".join(parts).strip()
                if not txt:
                    txt = f"[DOCX-no-visible-text: {os.path.basename(path)}]"
            except Exception:
                txt = f"[DOCX: {os.path.basename(path)}]"

        elif ext == ".csv":
            if pd is not None:
                try:
                    df = pd.read_csv(path)
                except Exception:
                    df = pd.read_csv(path, sep=None, engine="python")
                txt = df.head(200).to_csv(index=False)
            else:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    txt = f.read()

        elif ext == ".pdf":
            max_pages = 50 if not for_snippet else 2
            txt = _pdf_extract_text_pypdf(path, max_pages)
            if not txt:
                txt = _pdf_extract_text_pdfium(path, max_pages)
            if not txt:
                txt = _pdf_ocr_first_pages(path, pages_to_ocr=(2 if for_snippet else 3), scale=2.0)
            if not txt:
                txt = f"[PDF: {os.path.basename(path)}]"

    except Exception:
        txt = ""

    if for_snippet and len(txt) > 1200:
        txt = txt[:1200] + " …"
    return txt


def _build_bm25():
    global bm25, bm25_corpus
    bm25_corpus = []
    for m in meta:
        path = m.get("path")
        if not path:
            continue
        text = _extract_text(path, for_snippet=False)
        if not text:
            continue
        bm25_corpus.append(_tok(text))
    bm25 = BM25Okapi(bm25_corpus) if bm25_corpus else None


_load_rag()
_build_bm25()


def _load_snippet(m):
    path = m.get("path")
    if not path:
        return f"[file: {m.get('file','unknown')}]"
    snip = _extract_text(path, for_snippet=True)
    if not snip.strip():
        return f"[file: {os.path.basename(path)}]"
    return snip


def retrieve(query: str, k: int = TOP_K, mode: str = "dense"):
    
    results: list[tuple[float, dict, str]] = []

    if mode == "dense" and faiss_index is not None and embedder is not None and len(meta) > 0:
        qv = _embed([query])
        if qv is not None:
            qv = qv.astype("float32")
            D, I = faiss_index.search(qv, k)
            for dist, idx in zip(D[0], I[0]):
                if 0 <= idx < len(meta):
                    m = meta[idx]
                    snippet = _load_snippet(m)
                    results.append((float(-dist), m, snippet))  
        return results

    if mode == "bm25" and bm25 is not None and len(meta) > 0:
        toks = _tok(query)                      
        scores = bm25.get_scores(toks)
        topk = np.argsort(scores)[::-1][:k]
        for idx in topk:
            if 0 <= idx < len(meta):
                m = meta[idx]
                snippet = _load_snippet(m)
                results.append((float(scores[idx]), m, snippet))
        return results

    if mode == "hybrid":
       
        W_DENSE = 1.0   
        W_BM25  = 2.0   

        dense_hits = retrieve(query, k, mode="dense") if (faiss_index is not None and embedder is not None) else []
        bm_hits    = retrieve(query, k, mode="bm25")  if (bm25 is not None) else []

        
        all_hits: dict[str, tuple[float, dict, str]] = {}

        for s, m, snip in dense_hits:
            fid = m.get("file")
            all_hits[fid] = (all_hits.get(fid, (0.0, m, snip))[0] + W_DENSE * s, m, snip)

        for s, m, snip in bm_hits:
            fid = m.get("file")
            all_hits[fid] = (all_hits.get(fid, (0.0, m, snip))[0] + W_BM25 * s, m, snip)

        results = sorted(all_hits.values(), key=lambda x: -x[0])[:k]
        return results

    
    return results


def _rebuild_index():
    
    docs, meta_list = [], []
    files_indexed, files_skipped = [], []

    for root, _, files in os.walk(KNOW_DIR):
        for fn in files:
            path = os.path.join(root, fn)
            ext = os.path.splitext(fn)[1].lower()
            if ext not in SUPPORTED_EXTS:
                files_skipped.append({"file": fn, "reason": f"unsupported_ext:{ext}"})
                continue

            try:
                text = _extract_text(path, for_snippet=False)
            except Exception:
                text = ""
            
            if text and text.strip():
                docs.append(text)
                meta_list.append({"file": fn, "path": path})
                files_indexed.append(fn)
            else:
                files_skipped.append({"file": fn, "reason": "empty_extraction"})

    
    if not docs:
        if os.path.exists(INDEX_FILE): os.remove(INDEX_FILE)
        if os.path.exists(META_FILE):  os.remove(META_FILE)
        globals()["faiss_index"] = None
        globals()["meta"] = []
        globals()["bm25"] = None
        globals()["bm25_corpus"] = []
        return False, 0, files_indexed, files_skipped

   
    with open(META_FILE, "wb") as f:
        pickle.dump(meta_list, f)

   
    emb = _embed(docs)
    if emb is None:
        _load_rag()
        _build_bm25()
        return True, len(docs), files_indexed, files_skipped

   
    emb = emb.astype("float32")
    index = faiss.IndexFlatL2(emb.shape[1])
    index.add(emb)
    faiss.write_index(index, INDEX_FILE)

    
    _load_rag()
    _build_bm25()
    return True, len(docs), files_indexed, files_skipped


def build_prompt(user_query: str, hits, mode: str = "simple"):
    system = (
        "You are AxiomAI, a domain-specialist assistant for FORMAL VERIFICATION "
        "(safety & security included) and MATHEMATICAL PROOFS. Be precise, concise, and rigorous.\n"
        "• Always answer the latest user message from your own knowledge.\n"
        "• If the context below is useful, weave it in and cite like [1], [2].\n"
    )
    
    if mode == "simple":
        system += "Keep answers short, clear, and beginner-friendly.\n"
    elif mode == "academic":
        system += "Write in an academic style with depth and precision.\n"
    elif mode == "compare":
        system += "Focus on comparing and contrasting tools and methods.\n"
    elif mode == "references":
        system += "Prioritize citing the retrieved context explicitly in your answer.\n"
    elif mode == "steps":
        system += "Explain step by step, using numbered or bulleted points.\n"

    ctx_blocks = []
    for i, (_, m, snip) in enumerate(hits, 1):
        fname = m.get("file", f"doc{i}")
        ctx_blocks.append(f"[{i}] {fname}\n{snip}\n")
    context_text = "\n".join(ctx_blocks)
    if len(context_text) > MAX_CONTEXT_CHARS:
        context_text = context_text[:MAX_CONTEXT_CHARS] + "\n… (truncated)\n"

    user = (
        "Answer primarily from your own knowledge. If the context is relevant, use it and cite.\n\n"
        f"Context (optional):\n{context_text or '[no retrieved context]'}\n"
        f"Question: {user_query}"
    )
    return system, user


def _llm_generate(prompt: str, max_new_tokens: int = 300, temperature: float = 0.2) -> str:
    msgs = [
        {"role": "system", "content": "You generate **valid code** for formal verification tools. Return only code blocks with minimal comments."},
        {"role": "user", "content": prompt},
    ]
    input_ids = tokenizer.apply_chat_template(
        msgs, add_generation_prompt=True, return_tensors="pt"
    ).to(model.device)

    with torch.inference_mode():
        out = model.generate(
            input_ids=input_ids,
            max_new_tokens=max_new_tokens,
            do_sample=True,
            temperature=temperature,
            pad_token_id=tokenizer.eos_token_id,
            eos_token_id=tokenizer.eos_token_id,
        )
    gen_ids = out[0][input_ids.shape[-1]:]
    return tokenizer.decode(gen_ids, skip_special_tokens=True).strip()

def _strip_fences(text: str) -> str:
    fence = re.compile(r"^```[\w+-]*\s*|\s*```$", re.MULTILINE)
    return fence.sub("", text).strip()

def _prompt_for_tool(tool: str, goal: str, hints: str | None) -> str:
    hints_part = f"\nHints: {hints}" if hints else ""
    if tool.lower() == "coq":
        return (
            "Generate a minimal, compilable Coq proof script for the following goal.\n"
            "Use standard tactics; keep it short and correct.\n"
            f"Goal: {goal}{hints_part}\n"
            "Return only a code block."
        )
    if tool.lower() == "isabelle":
        return (
            "Generate an Isabelle/HOL theory snippet that proves the goal. Use 'theory ... imports Main begin' "
            "and end with 'end'. Keep it minimal and valid for Isabelle/jEdit.\n"
            f"Goal: {goal}{hints_part}\n"
            "Return only a code block."
        )
    return (
        "Generate a Python (z3py) script that models the constraints and prints solver result. "
        "Use from z3 import *; create a Solver(); assert constraints; then print(s.check()) "
        "and if sat print(model()).\n"
        f"Goal: {goal}{hints_part}\n"
        "Return only a code block."
    )

app = FastAPI(title="LLaMA-3 API (Stanage) + RAG + OCR + Logging + Upload + Proof API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   
    allow_methods=["*"],
    allow_headers=["*"],
)

class ChatIn(BaseModel):
    message: str
    history: list[dict] | None = None
    temperature: float | None = 0.2
    max_new_tokens: int | None = None
    use_rag: bool | None = True
    retrieval_mode: str | None = "hybrid"   
    mode: str | None = "simple"             

METRICS = {
    "start_time": time.time(),
    "requests": 0,
    "tokens_out": 0,
    "total_latency_sec": 0.0,
}
LOG_FILE = os.path.join(BASE_DIR, "chat_log.jsonl")

def log_event(obj: dict):
    obj["ts"] = datetime.utcnow().isoformat() + "Z"
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(json.dumps(obj, ensure_ascii=False) + "\n")

@app.get("/health")
def health():
    return {
        "status": "ok",
        "model": MODEL_ID,
        "device": str(device),
        "rag_index": bool(faiss_index is not None and len(meta) > 0),
        "rag_docs": len(meta),
        "bm25_ready": bool(bm25 is not None),
        "ocr": {
            "pdfium": bool(pdfium is not None),
            "pytesseract": bool(pytesseract is not None),
        },
        "z3_available": bool(z3 is not None),
    }

@app.get("/metrics")
def metrics():
    uptime = time.time() - METRICS["start_time"]
    avg_latency = (METRICS["total_latency_sec"] / METRICS["requests"]) if METRICS["requests"] else 0.0
    gpu = {}
    if torch.cuda.is_available():
        i = torch.cuda.current_device()
        gpu = {
            "name": torch.cuda.get_device_name(i),
            "memory_allocated_MB": int(torch.cuda.memory_allocated(i) / (1024 * 1024)),
            "memory_reserved_MB": int(torch.cuda.memory_reserved(i) / (1024 * 1024)),
        }
    return {
        "uptime_sec": int(uptime),
        "requests": METRICS["requests"],
        "avg_latency_sec": round(avg_latency, 3),
        "tokens_out": METRICS["tokens_out"],
        "device": str(device),
        "gpu": gpu,
    }

@app.post("/chat")
def chat(inp: ChatIn):
    t0 = time.time()

    query = (inp.message or "").strip()
    if not query:
        return {"reply": "Please type a question."}


    use_rag = True if inp.use_rag is None else bool(inp.use_rag)
    retrieval_mode = (inp.retrieval_mode or "hybrid").lower()
    if retrieval_mode not in ("dense", "bm25", "hybrid"):
        retrieval_mode = "hybrid"
    hits = retrieve(query, TOP_K, mode=retrieval_mode) if use_rag else []

 
    mode = (inp.mode or "simple").lower()
    MODE_DEFAULT_MAX = {
        "simple":     160,
        "academic":   700,
        "compare":    500,
        "references": 600,
        "steps":      450,
    }
    effective_max = int(inp.max_new_tokens) if inp.max_new_tokens else MODE_DEFAULT_MAX.get(mode, 400)
    effective_max = max(1, min(effective_max, 4000)) 

  
    system_msg, user_msg = build_prompt(query, hits, mode=mode)
    msgs = [{"role": "system", "content": system_msg},
            {"role": "user",   "content": user_msg}]

    input_ids = tokenizer.apply_chat_template(
        msgs, add_generation_prompt=True, return_tensors="pt"
    ).to(model.device)

    with torch.inference_mode():
        out = model.generate(
            input_ids=input_ids,
            max_new_tokens=effective_max,
            do_sample=True,
            temperature=(inp.temperature if inp.temperature is not None else 0.2),
            pad_token_id=tokenizer.eos_token_id,
            eos_token_id=tokenizer.eos_token_id,
        )

    gen_ids = out[0][input_ids.shape[-1]:]
    reply = tokenizer.decode(gen_ids, skip_special_tokens=True).strip()

   
    latency = time.time() - t0
    METRICS["requests"] += 1
    METRICS["total_latency_sec"] += latency
    METRICS["tokens_out"] += int(gen_ids.shape[0])

 
    log_event({
        "query": query,
        "reply": reply,
        "use_rag": use_rag,
        "retrieval_mode": retrieval_mode,     
        "mode": (inp.mode or "simple"),       
        "hits": [h[1].get("file") for h in hits],
        "latency_sec": round(latency, 3),
        "max_new_tokens": effective_max,
        "temperature": (inp.temperature if inp.temperature is not None else 0.2),
    })

    return {
        "reply": reply,
        "hits": [h[1].get("file") for h in hits],
        "retrieval_mode": retrieval_mode,
        "mode": mode,
    }

@app.post("/upload")
def upload(file: UploadFile = File(...)):
    fname = file.filename
    ext = os.path.splitext(fname)[1].lower()
    if ext not in SUPPORTED_EXTS:
        return {"ok": False, "error": f"Only {', '.join(sorted(SUPPORTED_EXTS))} supported."}

    dest = os.path.join(KNOW_DIR, fname)
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)

    ok, n, files_indexed, files_skipped = _rebuild_index()
    return {
        "ok": ok,
        "indexed_docs": n,
        "saved": fname,
        "files_indexed": files_indexed,
        "files_skipped": files_skipped
    }

@app.post("/reindex")
def reindex():
    ok, n, files_indexed, files_skipped = _rebuild_index()
    return {
        "ok": ok,
        "indexed_docs": n,
        "files_indexed": files_indexed,
        "files_skipped": files_skipped
    }


class ProofIn(BaseModel):
    goal: str
    tool: str | None = "coq"     
    hints: str | None = None
    max_new_tokens: int | None = 300
    temperature: float | None = 0.2

@app.post("/proof/generate")
def proof_generate(inp: ProofIn):
    tool = (inp.tool or "coq").lower()
    if tool not in ("coq", "isabelle", "z3py"):
        tool = "coq"
    prompt = _prompt_for_tool(tool, inp.goal, inp.hints)
    code = _llm_generate(prompt, max_new_tokens=inp.max_new_tokens or 300,
                         temperature=inp.temperature if inp.temperature is not None else 0.2)
    code = _strip_fences(code)
    return {"tool": tool, "code": code}

class Z3RunIn(BaseModel):
    goal: str
    hints: str | None = None
    timeout_sec: int | None = 5
    max_new_tokens: int | None = 250
    temperature: float | None = 0.2

def _exec_with_timeout(py_code: str, timeout_sec: int = 5):
    """
    Execute z3py code in a restricted env with a timeout.
    Returns (ran: bool, stdout: str, stderr: str, status: str)
    """
    if z3 is None:
        return False, "", "z3-solver not installed", "z3-missing"

    
    safe_builtins = {
        "print": print, "len": len, "range": range, "min": min, "max": max, "abs": abs, "str": str, "int": int, "float": float, "bool": bool
    }
    glob = {"__builtins__": safe_builtins, "z3": z3, "__name__": "__main__"}
    loc = {}

    stdout_buf, stderr_buf = io.StringIO(), io.StringIO()

    status = "ok"
    exc: Exception | None = None

    def run():
        nonlocal exc
        try:
            with contextlib.redirect_stdout(stdout_buf):
                with contextlib.redirect_stderr(stderr_buf):
                    exec(py_code, glob, loc)
        except Exception as e:
            exc = e

    th = threading.Thread(target=run, daemon=True)
    th.start()
    th.join(timeout=timeout_sec)
    if th.is_alive():
        status = "timeout"
        return False, stdout_buf.getvalue(), "execution timed out", status

    if exc is not None:
        status = "error"
        return False, stdout_buf.getvalue(), f"{type(exc).__name__}: {exc}", status

    return True, stdout_buf.getvalue(), stderr_buf.getvalue(), status

@app.post("/proof/z3/run")
def proof_z3_run(inp: Z3RunIn):
    prompt = _prompt_for_tool("z3py", inp.goal, inp.hints)
    code = _llm_generate(prompt, max_new_tokens=inp.max_new_tokens or 250,
                         temperature=inp.temperature if inp.temperature is not None else 0.2)
    code = _strip_fences(code)

    ran, out, err, status = _exec_with_timeout(code, timeout_sec=inp.timeout_sec or 5)
    return {
        "tool": "z3py",
        "code": code,
        "ran": ran,
        "status": status,
        "stdout": out,
        "stderr": err,
    }
# eval_report.py — summarize a run; add robust retrieval metrics; write CSV/XLSX + charts
import argparse, json, math, sys, os, io, json as pyjson
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

# -------- optional readers (best-effort; we degrade gracefully) --------
try:
    from docx import Document as _DocxDocument
except Exception:
    _DocxDocument = None

try:
    from pypdf import PdfReader as _PdfReader
except Exception:
    _PdfReader = None

# Optional semantic model (we require this now for retrieval quality;
# if missing, we degrade to NaN and warn once)
try:
    from sentence_transformers import SentenceTransformer, util as st_util
    _SEM_MODEL = SentenceTransformer("all-MiniLM-L6-v2")
except Exception:
    _SEM_MODEL = None


# ---------------- core loaders ----------------
def load_runs(run_dir: Path) -> pd.DataFrame:
    runs_path = run_dir / "runs.jsonl"
    if not runs_path.exists():
        raise FileNotFoundError(f"Could not find {runs_path}")

    rows = []
    with runs_path.open("r", encoding="utf-8") as f:
        for line in f:
            s = line.strip()
            if not s:
                continue
            rows.append(json.loads(s))

    df = pd.DataFrame(rows)

    # Backfill/normalize expected columns
    if "rag" not in df.columns and "use_rag" in df.columns:
        df["rag"] = df["use_rag"].map(lambda x: "ON" if x else "OFF")
    if "rag" not in df.columns:
        df["rag"] = "OFF"

    for col, default in [
        ("qid", None),
        ("question", ""),
        ("reply", ""),
        ("latency_sec", math.nan),
        ("tokens_out", math.nan),
        ("mode", df["mode"][0] if "mode" in df.columns and len(df) else "dense"),
        ("hits", []),
    ]:
        if col not in df.columns:
            df[col] = default

    # Types
    df["rag"] = df["rag"].astype(str).str.upper()
    df["latency_sec"] = pd.to_numeric(df["latency_sec"], errors="coerce")
    df["tokens_out"] = pd.to_numeric(df["tokens_out"], errors="coerce")

    # Handy view columns
    if "hits" in df.columns:
        df["hits_joined"] = df["hits"].apply(
            lambda h: ", ".join(map(str, h)) if isinstance(h, list) else str(h)
        )

    return df


def ensure_human_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Add empty columns you can score later (kept for manual use if desired)."""
    for col in ["accuracy", "completeness", "citations", "notes"]:
        if col not in df.columns:
            df[col] = "" if col == "notes" else pd.NA
    return df


# ---------------- simple auto metrics (kept from before) ----------------
def compute_keyword_hit(row, topk=3):
    q = str(row.get("question","")).lower()
    r = str(row.get("reply","")).lower()
    toks = [t for t in q.replace("/", " ").split() if t.isalpha()]
    toks = toks[:topk] if toks else []
    return int(any(t in r for t in toks)) if toks else np.nan


def compute_semantic_sim_answer(row):
    """Answer relevance (question vs answer) — requires _SEM_MODEL."""
    if _SEM_MODEL is None:
        return np.nan
    q = str(row.get("question",""))
    r = str(row.get("reply",""))
    if not q or not r:
        return np.nan
    qv = _SEM_MODEL.encode(q, convert_to_tensor=True, normalize_embeddings=True)
    rv = _SEM_MODEL.encode(r, convert_to_tensor=True, normalize_embeddings=True)
    sim = float(st_util.cos_sim(qv, rv).item())
    return sim


def add_basic_auto_metrics(df: pd.DataFrame) -> pd.DataFrame:
    df["kw_hit"]   = df.apply(compute_keyword_hit, axis=1)
    df["sem_sim"]  = df.apply(compute_semantic_sim_answer, axis=1)
    return df


# ---------------- retrieval content loader (best-effort) ----------------
SUPPORTED_EXTS = {".txt", ".md", ".docx", ".csv", ".pdf"}

def _read_text(path: Path, max_chars: int = 1500) -> str:
    """
    Best-effort text loader for evaluation time.
    We intentionally keep it lightweight and robust:
      - .txt/.md: utf-8 with ignore
      - .csv: pandas head(200)
      - .docx: python-docx paragraphs + tables (if available)
      - .pdf: pypdf text from first ~20 pages (if available)
    Returns text (possibly empty string) and truncates to max_chars.
    """
    if not path.exists() or not path.is_file():
        return ""

    ext = path.suffix.lower()

    try:
        if ext in (".txt", ".md"):
            with path.open("r", encoding="utf-8", errors="ignore") as f:
                txt = f.read()

        elif ext == ".csv":
            try:
                df = pd.read_csv(path)
            except Exception:
                df = pd.read_csv(path, sep=None, engine="python")
            txt = df.head(200).to_csv(index=False)

        elif ext == ".docx":
            if _DocxDocument is None:
                return f"[DOCX:{path.name}]"
            try:
                d = _DocxDocument(str(path))
                parts = []
                for p in d.paragraphs:
                    if p.text.strip():
                        parts.append(p.text.strip())
                for tbl in d.tables:
                    for row in tbl.rows:
                        row_txt = " | ".join((c.text or "").strip() for c in row.cells)
                        if row_txt.strip():
                            parts.append(row_txt)
                txt = "\n".join(parts).strip() or f"[DOCX:{path.name}]"
            except Exception:
                txt = f"[DOCX:{path.name}]"

        elif ext == ".pdf":
            if _PdfReader is None:
                return f"[PDF:{path.name}]"
            try:
                reader = _PdfReader(str(path))
                chunks = []
                # limit pages for speed
                for page in reader.pages[:20]:
                    chunks.append(page.extract_text() or "")
                txt = "\n".join(chunks).strip() or f"[PDF:{path.name}]"
            except Exception:
                txt = f"[PDF:{path.name}]"

        else:
            # Unsupported extension (shouldn't happen because we filter)
            txt = ""

    except Exception:
        txt = ""

    if max_chars and len(txt) > max_chars:
        txt = txt[:max_chars] + " …"
    return txt


# ---------------- semantic retrieval quality ----------------
def compute_semantic_retrieval_scores(row, knowledge_dir: Path) -> tuple[float, float, str]:
    """
    For each retrieved filename in 'hits', open the corresponding file in knowledge_dir,
    read a text snippet, compute embedding similarity with the question, and return:
      (max_similarity, avg_similarity, json_list_of_per_doc_scores)
    If nothing can be scored, returns (nan, nan, "[]").
    """
    if _SEM_MODEL is None:
        return np.nan, np.nan, "[]"

    hits = row.get("hits", [])
    if not isinstance(hits, list) or not len(hits):
        return np.nan, np.nan, "[]"

    q = str(row.get("question","")).strip()
    if not q:
        return np.nan, np.nan, "[]"

    sims = []
    qv = _SEM_MODEL.encode(q, convert_to_tensor=True, normalize_embeddings=True)

    for fname in hits:
        if not isinstance(fname, str):
            continue
        # We only allow supported extensions
        ext = Path(fname).suffix.lower()
        if ext and ext not in SUPPORTED_EXTS:
            continue

        fpath = knowledge_dir / fname
        if not fpath.exists():
            # Try to find case-insensitively (mac devs sometimes hit this)
            matches = list(knowledge_dir.glob(fname)) or \
                      [p for p in knowledge_dir.glob("*") if p.name.lower() == fname.lower()]
            if matches:
                fpath = matches[0]
            else:
                continue

        content = _read_text(fpath, max_chars=1500)
        if not content or content.startswith("[PDF:") or content.startswith("[DOCX:"):
            # still compute; these markers are fine but likely low-signal
            pass

        cv = _SEM_MODEL.encode(content, convert_to_tensor=True, normalize_embeddings=True)
        sim = float(st_util.cos_sim(qv, cv).item())
        sims.append({"file": fpath.name, "sim": sim})

    if not sims:
        return np.nan, np.nan, "[]"

    max_sim = max(s["sim"] for s in sims)
    avg_sim = sum(s["sim"] for s in sims) / len(sims)
    return max_sim, avg_sim, pyjson.dumps(sims, ensure_ascii=False)


def add_semantic_retrieval_metrics(df: pd.DataFrame, knowledge_dir: Path) -> pd.DataFrame:
    max_list, avg_list, per_doc_json = [], [], []
    for _, row in df.iterrows():
        m, a, js = compute_semantic_retrieval_scores(row, knowledge_dir)
        max_list.append(m); avg_list.append(a); per_doc_json.append(js)
    df["sem_retrieval_max"] = max_list
    df["sem_retrieval_avg"] = avg_list
    df["sem_retrieval_topk"] = per_doc_json
    return df


# ---------------- outputs ----------------
def write_tables(df: pd.DataFrame, run_dir: Path) -> tuple[Path, Path | None]:
    out_csv = run_dir / "report.csv"
    out_xlsx = run_dir / "report.xlsx"
    df.to_csv(out_csv, index=False)
    wrote_xlsx = None
    try:
        df.to_excel(out_xlsx, index=False)  # needs openpyxl or xlsxwriter
        wrote_xlsx = out_xlsx
    except Exception as e:
        print(f"(!) Could not write Excel: {e}. CSV is still saved.")
    return out_csv, wrote_xlsx


def bar_plot(avg_off: float, avg_on: float, title: str, ylabel: str, out_path: Path):
    plt.figure(figsize=(4.6, 3.6), dpi=150)
    xs = ["RAG OFF", "RAG ON"]
    ys = [avg_off, avg_on]
    plt.bar(xs, ys)
    plt.title(title)
    plt.ylabel(ylabel)
    plt.tight_layout()
    plt.savefig(out_path)
    plt.close()


# ---------------- run selection ----------------
def latest_run_dir(base: Path = Path("eval_runs")) -> Path | None:
    if not base.exists():
        return None
    dirs = [p for p in base.iterdir() if p.is_dir()]
    if not dirs:
        return None
    return max(dirs, key=lambda p: p.stat().st_mtime)


# ---------------- main ----------------
def main():
    ap = argparse.ArgumentParser(
        description="Summarize an evaluation run; compute semantic retrieval quality; save CSV/XLSX + charts."
    )
    ap.add_argument(
        "--run_dir",
        default=None,
        help="Path to eval_runs/<timestamp>. If omitted, use newest in eval_runs/."
    )
    ap.add_argument(
        "--knowledge_dir",
        default="knowledge",
        help="Path to your backend knowledge folder (where documents live)."
    )
    args = ap.parse_args()

    # Resolve dirs
    if args.run_dir:
        run_dir = Path(args.run_dir).expanduser().resolve()
    else:
        run_dir = latest_run_dir()
        if run_dir is None:
            print("No runs found in eval_runs/. Run eval_client.py first.", file=sys.stderr)
            sys.exit(1)

    knowledge_dir = Path(args.knowledge_dir).expanduser().resolve()
    if not knowledge_dir.exists():
        print(f"(!) knowledge_dir not found: {knowledge_dir}. Semantic retrieval will likely be NaN.", file=sys.stderr)

    df = load_runs(run_dir)
    df = ensure_human_columns(df)
    df = add_basic_auto_metrics(df)
    df = add_semantic_retrieval_metrics(df, knowledge_dir)

    # Averages by RAG condition
    avg = df.groupby("rag", dropna=False).agg(
        avg_latency_sec=("latency_sec", "mean"),
        avg_tokens=("tokens_out", "mean"),
        n=("qid", "count"),
        kw_hit_rate=("kw_hit", "mean"),
        sem_sim_avg=("sem_sim", "mean"),  # answer relevance
        # NEW: retrieval quality (semantic, uses document content)
        sem_retr_max=("sem_retrieval_max", "mean"),
        sem_retr_avg=("sem_retrieval_avg", "mean"),
    )

    # Extract a few for the charts
    avg_off = float(avg.loc["OFF", "avg_latency_sec"]) if "OFF" in avg.index else float("nan")
    avg_on  = float(avg.loc["ON",  "avg_latency_sec"]) if "ON"  in avg.index else float("nan")
    tok_off = float(avg.loc["OFF", "avg_tokens"]) if "OFF" in avg.index else float("nan")
    tok_on  = float(avg.loc["ON",  "avg_tokens"]) if "ON"  in avg.index else float("nan")

    # Write tables
    out_csv, out_xlsx = write_tables(df, run_dir)

    # Charts
    lat_png = run_dir / "latency_bar.png"
    tok_png = run_dir / "tokens_bar.png"
    bar_plot(avg_off, avg_on, "Average Latency", "seconds", lat_png)
    bar_plot(tok_off, tok_on, "Average Tokens", "tokens", tok_png)

    # Console summary
    print("\n=== Evaluation Report ===")
    print(f"Run dir:   {run_dir}")
    print(f"Rows:      {len(df)}")
    print(f"CSV:       {out_csv}")
    print(f"Excel:     {out_xlsx if out_xlsx else '(not written)'}")
    print("\nAverages:")
    print(f"  Avg latency (RAG OFF): {avg_off:.3f} s")
    print(f"  Avg latency (RAG ON):  {avg_on:.3f} s")
    print(f"  Avg tokens  (RAG OFF): {tok_off:.1f}")
    print(f"  Avg tokens  (RAG ON):  {tok_on:.1f}")
    print("\nContent/Retrieval metrics (NaN means not available):")
    shown_cols = ["kw_hit_rate","sem_sim_avg","sem_retr_max","sem_retr_avg"]
    print(avg[shown_cols])
    print("\nSaved charts:")
    print(f"  {lat_png.name}, {tok_png.name}")
    print("\nTip: Open the CSV/Excel to inspect per-question 'sem_retrieval_topk' (per-doc sims).")
    if "qid_custom" in df.columns:
        print("Note: Custom IDs available in column 'qid_custom'.")

if __name__ == "__main__":
    main()